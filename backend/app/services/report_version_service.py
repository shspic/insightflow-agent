import hashlib
import html
import json
import re
import shutil
import uuid
from datetime import datetime
from app.core.timeutils import utcnow
from pathlib import Path
from typing import Any, Literal

from sqlalchemy import func, select
from sqlalchemy.orm import Session

from app.core.config import BACKEND_DIR, settings
from app.models.report import Report
from app.models.report_asset import ReportAsset
from app.models.task import Task
from app.services.report_template_service import get_report_template
from app.services.workspace_service import safe_public_text


ExportFormat = Literal["markdown", "docx", "pdf"]
_MIME_TYPES = {
    "markdown": "text/markdown; charset=utf-8",
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "pdf": "application/pdf",
}


class ReportVersionError(Exception):
    def __init__(self, message: str, code: str = "REPORT_ERROR") -> None:
        self.message = message
        self.code = code
        super().__init__(message)


def create_report_version(
    db: Session,
    *,
    task: Task,
    state: dict[str, Any],
    template_key: str | None = None,
    generation_source: str = "initial",
    correction_note: str | None = None,
) -> Report:
    if task.workspace_id is None or task.owner_user_id is None:
        raise ReportVersionError("V2 报告必须关联工作区和所有者", "REPORT_SCOPE_REQUIRED")
    template = get_report_template(template_key or _task_template(task))
    existing_count = db.scalar(
        select(func.count(Report.id)).where(Report.task_id == task.id)
    ) or 0
    if existing_count >= settings.report_history_max_versions:
        raise ReportVersionError("已达到单任务报告历史版本上限", "REPORT_VERSION_LIMIT")
    next_version = (
        db.scalar(select(func.max(Report.version)).where(Report.task_id == task.id)) or 0
    ) + 1
    content = _build_markdown(
        task=task,
        state=state,
        template_key=template.template_key,
        correction_note=correction_note,
    )
    content_hash = hashlib.sha256(content.encode("utf-8")).hexdigest()
    if generation_source == "retry":
        reused = db.scalar(
            select(Report).where(
                Report.task_id == task.id,
                Report.content_hash == content_hash,
                Report.status.in_(["ready", "ready_with_warnings"]),
            )
        )
        if reused is not None:
            task.report_id = reused.id
            task.report_path = _compatibility_report_path(reused)
            db.flush()
            return reused

    quality = _quality_summary(state, template.required_sections, content)
    warnings = list(dict.fromkeys(str(item)[:500] for item in (state.get("warnings") or [])))
    status = "ready_with_warnings" if warnings or quality["status"] != "passed" else "ready"
    now = utcnow()
    report = Report(
        task_id=task.id,
        workspace_id=task.workspace_id,
        owner_user_id=task.owner_user_id,
        version=next_version,
        status=status,
        title=f"{template.display_name}（版本 {next_version}）",
        template_key=template.template_key,
        language="zh-CN",
        markdown_content=content,
        content_hash=content_hash,
        generation_source=generation_source,
        quality_status=quality["status"],
        quality_summary_json=json.dumps(quality, ensure_ascii=False),
        warnings_json=json.dumps(warnings, ensure_ascii=False),
        completed_at=now,
    )
    db.add(report)
    db.flush()
    _persist_text_asset(db, report=report, content=content, asset_type="markdown", suffix="md")
    _persist_citation_manifest(db, report=report, state=state)
    _copy_chart_assets(db, report=report, state=state)

    previous = db.scalars(
        select(Report).where(
            Report.task_id == task.id,
            Report.id != report.id,
            Report.status.in_(["ready", "ready_with_warnings"]),
        )
    ).all()
    for item in previous:
        item.superseded_at = now
    task.report_id = report.id
    task.report_path = _compatibility_report_path(report)
    db.flush()
    return report


def list_owned_reports(
    db: Session, *, workspace_id: int, task_id: int, owner_user_id: int
) -> list[Report]:
    return list(
        db.scalars(
            select(Report)
            .where(
                Report.workspace_id == workspace_id,
                Report.task_id == task_id,
                Report.owner_user_id == owner_user_id,
            )
            .order_by(Report.version.desc())
        ).all()
    )


def get_owned_report(
    db: Session,
    *,
    workspace_id: int,
    task_id: int,
    report_id: int,
    owner_user_id: int,
) -> Report:
    report = db.scalar(
        select(Report).where(
            Report.id == report_id,
            Report.workspace_id == workspace_id,
            Report.task_id == task_id,
            Report.owner_user_id == owner_user_id,
        )
    )
    if report is None:
        raise ReportVersionError("报告不存在或无权访问", "REPORT_NOT_FOUND")
    return report


def export_report(db: Session, *, report: Report, export_format: ExportFormat) -> ReportAsset:
    if export_format not in _MIME_TYPES:
        raise ReportVersionError("不支持的报告导出格式", "UNSUPPORTED_EXPORT_FORMAT")
    existing = db.scalar(
        select(ReportAsset).where(
            ReportAsset.report_id == report.id,
            ReportAsset.asset_type == export_format,
            ReportAsset.checksum == report.content_hash,
            ReportAsset.status == "ready",
            ReportAsset.deleted_at.is_(None),
        )
    )
    if existing is not None and resolve_asset_path(existing).is_file():
        return existing

    suffix = "md" if export_format == "markdown" else export_format
    asset = _new_asset(
        report=report,
        asset_type=export_format,
        suffix=suffix,
        mime_type=_MIME_TYPES[export_format],
        checksum=report.content_hash,
    )
    db.add(asset)
    db.flush()
    path = resolve_asset_path(asset)
    path.parent.mkdir(parents=True, exist_ok=True)
    try:
        if export_format == "markdown":
            path.write_text(report.markdown_content, encoding="utf-8")
        elif export_format == "docx":
            _write_docx(report, path, _report_assets(db, report.id))
        else:
            _write_pdf(report, path, _report_assets(db, report.id))
        asset.size_bytes = path.stat().st_size
        asset.status = "ready"
        db.flush()
        return asset
    except Exception as exc:
        asset.status = "failed"
        asset.metadata_json = json.dumps({"error": safe_public_text(str(exc))[:500]}, ensure_ascii=False)
        if path.is_file():
            path.unlink()
        db.flush()
        raise ReportVersionError(f"报告导出失败：{safe_public_text(str(exc))}", "REPORT_EXPORT_FAILED") from exc


def delete_report_version(db: Session, *, task: Task, report: Report) -> None:
    usable = list(
        db.scalars(
            select(Report).where(
                Report.task_id == task.id,
                Report.status.in_(["ready", "ready_with_warnings"]),
            )
        ).all()
    )
    if len(usable) <= 1 or task.report_id == report.id:
        raise ReportVersionError("不能删除当前版本或任务唯一可用报告", "REPORT_DELETE_BLOCKED")
    now = utcnow()
    report.status = "superseded"
    report.superseded_at = now
    for asset in _report_assets(db, report.id):
        asset.status = "superseded"
    db.flush()


def set_current_report(db: Session, *, task: Task, report: Report) -> Report:
    if report.task_id != task.id or report.owner_user_id != task.owner_user_id:
        raise ReportVersionError("报告不属于当前任务", "REPORT_SCOPE_MISMATCH")
    if report.status not in {"ready", "ready_with_warnings"}:
        raise ReportVersionError("只有可用报告版本可以设为当前版本", "REPORT_NOT_READY")
    now = utcnow()
    for item in db.scalars(
        select(Report).where(
            Report.task_id == task.id,
            Report.id != report.id,
            Report.status.in_(["ready", "ready_with_warnings"]),
        )
    ).all():
        item.superseded_at = item.superseded_at or now
    report.superseded_at = None
    task.report_id = report.id
    task.report_path = _compatibility_report_path(report)
    db.flush()
    return report


def resolve_asset_path(asset: ReportAsset) -> Path:
    key_path = Path(asset.storage_key)
    if key_path.is_absolute() or ".." in key_path.parts:
        raise ReportVersionError("报告资产标识无效", "INVALID_STORAGE_KEY")
    root = _storage_root().resolve()
    resolved = (root / key_path).resolve()
    if root != resolved and root not in resolved.parents:
        raise ReportVersionError("报告资产超出存储目录", "INVALID_STORAGE_KEY")
    return resolved


def report_response(db: Session, *, task: Task, report: Report) -> dict[str, Any]:
    assets = []
    for item in _report_assets(db, report.id):
        assets.append(
            {
                "id": item.id,
                "asset_type": item.asset_type,
                "format": item.format,
                "display_name": item.display_name,
                "mime_type": item.mime_type,
                "size_bytes": item.size_bytes,
                "checksum": item.checksum,
                "status": item.status,
                "download_url": (
                    f"/api/v2/workspaces/{report.workspace_id}/tasks/{report.task_id}"
                    f"/reports/{report.id}/assets/{item.id}/download"
                    if item.status == "ready" and item.deleted_at is None
                    else None
                ),
                "created_at": item.created_at,
            }
        )
    return {
        "id": report.id,
        "task_id": report.task_id,
        "workspace_id": report.workspace_id,
        "version": report.version,
        "status": report.status,
        "title": report.title,
        "template_key": report.template_key,
        "language": report.language,
        "markdown_content": report.markdown_content,
        "generation_source": report.generation_source,
        "quality_status": report.quality_status,
        "quality_summary": _json_dict(report.quality_summary_json),
        "warnings": _json_list(report.warnings_json),
        "is_current": task.report_id == report.id,
        "assets": assets,
        "created_at": report.created_at,
        "completed_at": report.completed_at,
    }


def _build_markdown(
    *, task: Task, state: dict[str, Any], template_key: str, correction_note: str | None
) -> str:
    template = get_report_template(template_key)
    context = state.get("workspace_context") or {}
    findings = state.get("analysis_findings") or []
    evidence = state.get("document_evidence") or []
    charts = state.get("chart_assets") or []
    warnings = state.get("warnings") or []
    assumptions = state.get("assumptions") or []
    facts = "\n".join(f"- {json.dumps(item, ensure_ascii=False, default=str)}" for item in findings)
    citations = []
    for index, item in enumerate(evidence, start=1):
        location = (
            f"第 {item.get('page_number')} 页"
            if item.get("page_number") is not None
            else item.get("section_path") or "位置未标注"
        )
        citations.append(
            f"- [{index}] {safe_public_text(str(item.get('filename') or '来源'))}，"
            f"{safe_public_text(str(location))}：{safe_public_text(str(item.get('snippet') or ''))}"
        )
    section_content = {
        "执行摘要": "本报告基于已完成的数据分析与文档检索结果生成，不在报告阶段重新计算数据。",
        "摘要": "本报告围绕既定研究问题整理已有分析结果和可追溯证据。",
        "匹配摘要": "本报告比较岗位要求与候选人材料中已有的可验证证据。",
        "任务与资料范围": _file_scope(context),
        "研究问题": safe_public_text(state.get("clarified_request") or task.user_input),
        "岗位要求": "\n".join(citations) or "当前资料中未提取到明确岗位要求。",
        "资料与方法": _methods(findings),
        "分析方法": _methods(findings),
        "分析结果": facts or "当前没有可复核的结构化分析结论。",
        "关键发现": facts or "当前没有可复核的结构化分析结论。",
        "候选人证据": "\n".join(citations) or "当前没有可定位的候选人证据。",
        "讨论": _synthesis(findings, evidence),
        "优势与差距": _synthesis(findings, evidence),
        "数据质量与异常风险": _bullets(
            [*context.get("data_quality_issues", []), *warnings]
            or ["未记录额外异常；关键结论仍建议人工复核。"]
        ),
        "风险与待确认事项": _bullets(warnings or ["重要判断仍应由用户结合原始材料复核。"]),
        "图表与表格": _chart_markdown(charts),
        "引用与证据": "\n".join(citations) or "当前未检索到可定位引用。",
        "结论与建议": _recommendations(findings, evidence, warnings),
        "行动建议": _recommendations(findings, evidence, warnings),
        "假设与限制": _bullets(
            assumptions
            + [
                "仅使用当前任务已确认的文件和已完成步骤结果。",
                "OCR、文档检索、采样统计和模型组织均可能存在误差。",
            ]
        ),
        "Quality Review": _quality_markdown(state),
    }
    lines = [f"# {template.display_name}", "", f"> 任务 ID：{task.id}", ""]
    if correction_note:
        lines.extend(["## 本次纠正说明", safe_public_text(correction_note) or "", ""])
    for section in template.required_sections:
        lines.extend([f"## {section}", section_content.get(section, "本节暂无补充内容。"), ""])
    for section in template.optional_sections:
        if section == "多文件综合结论":
            lines.extend([f"## {section}", _synthesis(findings, evidence), ""])
    return "\n".join(lines).strip() + "\n"


def _quality_summary(
    state: dict[str, Any], required_sections: tuple[str, ...], content: str
) -> dict[str, Any]:
    missing = [section for section in required_sections if f"## {section}" not in content]
    review = state.get("quality_review") or state.get("review_result") or {}
    review_status = str(review.get("status") or review.get("quality_status") or "")
    status = "failed" if missing or review_status in {"failed", "blocked"} else "passed"
    return {
        "status": status,
        "missing_sections": missing,
        "source_review_status": review_status or "deterministic_check",
    }


def _quality_markdown(state: dict[str, Any]) -> str:
    review = state.get("quality_review") or state.get("review_result") or {}
    if not review:
        return "- 状态：passed（已通过必需章节与结构一致性检查）"
    return f"- 状态：{safe_public_text(str(review.get('status') or 'completed'))}\n- 摘要：{safe_public_text(json.dumps(review, ensure_ascii=False, default=str)[:1500])}"


def _file_scope(context: dict[str, Any]) -> str:
    files = context.get("files") or []
    if not files:
        return "本任务未记录文件范围。"
    return "\n".join(
        f"- {safe_public_text(str(item.get('filename') or '未命名文件'))}；"
        f"角色：{safe_public_text(str(item.get('effective_role') or '未确认'))}"
        for item in files
    )


def _methods(findings: list[dict[str, Any]]) -> str:
    methods = [item.get("calculation_method") for item in findings if item.get("calculation_method")]
    return _bullets(methods or ["读取文件 Profile 与 Workspace Context", "执行受限的预设分析或检索工具"])


def _bullets(items: list[Any]) -> str:
    lines = []
    for item in items:
        text = json.dumps(item, ensure_ascii=False, default=str) if isinstance(item, (dict, list)) else str(item)
        lines.append(f"- {safe_public_text(text)}")
    return "\n".join(lines)


def _chart_markdown(charts: list[dict[str, Any]]) -> str:
    active = [item for item in charts if not item.get("skipped") and item.get("asset_name")]
    if not active:
        return "本任务未生成可用图表。"
    return "\n".join(
        f"- {safe_public_text(str(item.get('title') or '图表'))}："
        f"{safe_public_text(str(Path(str(item['asset_name'])).name))}"
        for item in active
    )


def _synthesis(findings: list[dict[str, Any]], evidence: list[dict[str, Any]]) -> str:
    if findings and evidence:
        return "已并列呈现结构化数据结论和文档证据；未确认连接字段时不做行级拼接。"
    if findings:
        return "当前综合结论主要基于结构化数据分析。"
    if evidence:
        return "当前综合结论主要基于文档检索证据。"
    return "当前证据不足，无法形成可靠的综合结论。"


def _recommendations(
    findings: list[dict[str, Any]], evidence: list[dict[str, Any]], warnings: list[Any]
) -> str:
    items = ["优先复核关键数字、异常值和引用定位。"]
    if findings and evidence:
        items.append("如需行级关联，请先确认连接字段与文件关系。")
    if warnings:
        items.append("处理 Quality Review 警告后再用于正式决策。")
    return _bullets(items)


def _task_template(task: Task) -> str:
    try:
        preferences = json.loads(task.report_preferences_json or "{}")
    except json.JSONDecodeError:
        preferences = {}
    return str(preferences.get("template_key") or "comprehensive_analysis")


def _storage_root() -> Path:
    report_dir = Path(settings.report_dir)
    if not report_dir.is_absolute():
        report_dir = BACKEND_DIR / report_dir
    return report_dir.parent


def _asset_key(report: Report, suffix: str) -> str:
    token = uuid.uuid4().hex
    return (
        f"reports/assets/u{report.owner_user_id}/w{report.workspace_id}/"
        f"t{report.task_id}/r{report.id}/{token}.{suffix}"
    )


def _new_asset(
    *,
    report: Report,
    asset_type: str,
    suffix: str,
    mime_type: str,
    checksum: str | None = None,
    display_name: str | None = None,
) -> ReportAsset:
    return ReportAsset(
        report_id=report.id,
        task_id=report.task_id,
        workspace_id=report.workspace_id,
        owner_user_id=report.owner_user_id,
        asset_type=asset_type,
        format=suffix,
        display_name=display_name or f"insightflow-report-v{report.version}.{suffix}",
        storage_key=_asset_key(report, suffix),
        mime_type=mime_type,
        size_bytes=0,
        checksum=checksum,
        status="generating",
    )


def _persist_text_asset(
    db: Session, *, report: Report, content: str, asset_type: str, suffix: str
) -> ReportAsset:
    checksum = hashlib.sha256(content.encode("utf-8")).hexdigest()
    asset = _new_asset(
        report=report,
        asset_type=asset_type,
        suffix=suffix,
        mime_type=_MIME_TYPES["markdown"],
        checksum=checksum,
    )
    path = resolve_asset_path(asset)
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(content, encoding="utf-8")
    asset.size_bytes = path.stat().st_size
    asset.status = "ready"
    db.add(asset)
    return asset


def _persist_citation_manifest(db: Session, *, report: Report, state: dict[str, Any]) -> None:
    content = json.dumps(state.get("document_evidence") or [], ensure_ascii=False, default=str)
    asset = _new_asset(
        report=report,
        asset_type="citation_manifest",
        suffix="json",
        mime_type="application/json",
        checksum=hashlib.sha256(content.encode("utf-8")).hexdigest(),
        display_name=f"citation-manifest-v{report.version}.json",
    )
    path = resolve_asset_path(asset)
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(content, encoding="utf-8")
    asset.size_bytes = path.stat().st_size
    asset.status = "ready"
    db.add(asset)


def _copy_chart_assets(db: Session, *, report: Report, state: dict[str, Any]) -> None:
    chart_root = Path(settings.chart_dir)
    if not chart_root.is_absolute():
        chart_root = BACKEND_DIR / chart_root
    chart_root = chart_root.resolve()
    for item in state.get("chart_assets") or []:
        name = Path(str(item.get("asset_name") or "")).name
        if not name:
            continue
        source = (chart_root / name).resolve()
        if chart_root not in source.parents or not source.is_file():
            continue
        suffix = source.suffix.lower().removeprefix(".") or "png"
        asset = _new_asset(
            report=report,
            asset_type="chart",
            suffix=suffix,
            mime_type="image/png" if suffix == "png" else "image/jpeg",
            display_name=name,
        )
        destination = resolve_asset_path(asset)
        destination.parent.mkdir(parents=True, exist_ok=True)
        shutil.copyfile(source, destination)
        asset.size_bytes = destination.stat().st_size
        asset.checksum = _file_hash(destination)
        asset.status = "ready"
        db.add(asset)


def _report_assets(db: Session, report_id: int) -> list[ReportAsset]:
    return list(
        db.scalars(
            select(ReportAsset)
            .where(ReportAsset.report_id == report_id, ReportAsset.deleted_at.is_(None))
            .order_by(ReportAsset.created_at.asc())
        ).all()
    )


def _write_docx(report: Report, path: Path, assets: list[ReportAsset]) -> None:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt

    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    styles = document.styles
    styles["Normal"].font.name = "Microsoft YaHei"
    styles["Normal"].font.size = Pt(10.5)
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    for style_name in ("Title", "Heading 1", "Heading 2", "Heading 3"):
        styles[style_name].font.name = "Microsoft YaHei"
        styles[style_name]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run(
        f"InsightFlow Agent · 报告版本 {report.version} · "
        f"{(report.completed_at or report.created_at).strftime('%Y-%m-%d %H:%M UTC')}"
    )
    _markdown_to_docx(document, report.markdown_content)
    chart_assets = [item for item in assets if item.asset_type == "chart" and item.status == "ready"]
    if chart_assets:
        document.add_heading("报告图表资产", level=1)
        for asset in chart_assets:
            document.add_paragraph(asset.display_name)
            chart_path = resolve_asset_path(asset)
            if chart_path.is_file():
                try:
                    document.add_picture(str(chart_path), width=Inches(6.5))
                except Exception:
                    document.add_paragraph("图表暂无法嵌入，请通过报告资产下载。")
            else:
                document.add_paragraph("图表文件不存在。")
    document.save(path)


def _markdown_to_docx(document: Any, markdown: str) -> None:
    lines = markdown.splitlines()
    index = 0
    while index < len(lines):
        line = lines[index].strip()
        if not line:
            index += 1
            continue
        if line.startswith("```") or line.startswith("<"):
            index += 1
            continue
        if line.startswith("#"):
            level = min(3, len(line) - len(line.lstrip("#")))
            document.add_heading(_plain(line[level:].strip()), level=level - 1 if level > 1 else 0)
            index += 1
            continue
        if (
            "|" in line
            and index + 1 < len(lines)
            and re.match(r"^\|?\s*:?-{3,}", lines[index + 1].strip())
        ):
            rows = [_table_cells(line)]
            index += 2
            while index < len(lines) and "|" in lines[index]:
                rows.append(_table_cells(lines[index]))
                index += 1
            width = max(len(row) for row in rows)
            table = document.add_table(rows=len(rows), cols=width)
            table.style = "Table Grid"
            for row_index, row in enumerate(rows):
                for column_index, value in enumerate(row):
                    table.cell(row_index, column_index).text = _plain(value)
            continue
        if line.startswith(("- ", "* ")):
            document.add_paragraph(_plain(line[2:]), style="List Bullet")
        elif re.match(r"^\d+[.)]\s+", line):
            document.add_paragraph(_plain(re.sub(r"^\d+[.)]\s+", "", line)), style="List Number")
        else:
            document.add_paragraph(_plain(line.lstrip("> ")))
        index += 1


def _write_pdf(report: Report, path: Path, assets: list[ReportAsset]) -> None:
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import mm
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.platypus import Image, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

    font_name = _register_pdf_font(pdfmetrics, TTFont)
    styles = getSampleStyleSheet()
    normal = ParagraphStyle(
        "InsightFlowNormal", parent=styles["BodyText"], fontName=font_name, fontSize=9.5, leading=15
    )
    headings = {
        1: ParagraphStyle(
            "InsightFlowTitle", parent=styles["Title"], fontName=font_name, fontSize=20, leading=26,
            alignment=TA_CENTER, spaceAfter=10,
        ),
        2: ParagraphStyle(
            "InsightFlowH2", parent=styles["Heading2"], fontName=font_name, fontSize=14, leading=20,
            textColor=colors.HexColor("#17365D"), spaceBefore=9, spaceAfter=5,
        ),
        3: ParagraphStyle(
            "InsightFlowH3", parent=styles["Heading3"], fontName=font_name, fontSize=11, leading=16,
            spaceBefore=6, spaceAfter=3,
        ),
    }
    story: list[Any] = []
    lines = report.markdown_content.splitlines()
    index = 0
    while index < len(lines):
        line = lines[index].strip()
        if not line or line.startswith("```") or line.startswith("<"):
            index += 1
            continue
        if line.startswith("#"):
            level = min(3, len(line) - len(line.lstrip("#")))
            story.append(Paragraph(_pdf_text(_plain(line[level:].strip())), headings[level]))
        elif (
            "|" in line
            and index + 1 < len(lines)
            and re.match(r"^\|?\s*:?-{3,}", lines[index + 1].strip())
        ):
            rows = [_table_cells(line)]
            index += 2
            while index < len(lines) and "|" in lines[index]:
                rows.append(_table_cells(lines[index]))
                index += 1
            data = [[Paragraph(_pdf_text(_plain(cell)), normal) for cell in row] for row in rows]
            table = Table(data, repeatRows=1)
            table.setStyle(
                TableStyle(
                    [
                        ("FONTNAME", (0, 0), (-1, -1), font_name),
                        ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#AAB7C4")),
                        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#EAF0F6")),
                        ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ]
                )
            )
            story.extend([table, Spacer(1, 4 * mm)])
            continue
        else:
            prefix = "• " if line.startswith(("- ", "* ")) else ""
            text = line[2:] if prefix else line.lstrip("> ")
            story.append(Paragraph(_pdf_text(prefix + _plain(text)), normal))
            story.append(Spacer(1, 1.5 * mm))
        index += 1
    charts = [item for item in assets if item.asset_type == "chart" and item.status == "ready"]
    if charts:
        story.append(Paragraph("报告图表资产", headings[2]))
        for asset in charts:
            chart_path = resolve_asset_path(asset)
            story.append(Paragraph(_pdf_text(asset.display_name), normal))
            if chart_path.is_file():
                try:
                    image = Image(str(chart_path))
                    image._restrictSize(170 * mm, 95 * mm)
                    story.append(image)
                except Exception:
                    story.append(Paragraph("图表暂无法嵌入，请通过报告资产下载。", normal))
    document = SimpleDocTemplate(
        str(path), pagesize=A4, leftMargin=18 * mm, rightMargin=18 * mm,
        topMargin=18 * mm, bottomMargin=18 * mm,
        title=report.title, author="InsightFlow Agent",
    )
    document.build(story, onFirstPage=_pdf_footer(report, font_name), onLaterPages=_pdf_footer(report, font_name))


def _register_pdf_font(pdfmetrics: Any, ttfont: Any) -> str:
    candidates = [
        Path("C:/Windows/Fonts/msyh.ttc"),
        Path("C:/Windows/Fonts/simsun.ttc"),
        Path("/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc"),
        Path("/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc"),
    ]
    for candidate in candidates:
        if not candidate.is_file():
            continue
        try:
            pdfmetrics.registerFont(ttfont("InsightFlowCJK", str(candidate), subfontIndex=0))
            return "InsightFlowCJK"
        except Exception:
            continue
    raise RuntimeError("未找到可用中文字体；请安装 Microsoft YaHei、SimSun 或 Noto Sans CJK")


def _pdf_footer(report: Report, font_name: str):
    def draw(canvas: Any, document: Any) -> None:
        canvas.saveState()
        canvas.setFont(font_name, 8)
        canvas.drawString(18 * 2.8346, 12 * 2.8346, f"InsightFlow Agent · v{report.version}")
        canvas.drawRightString(192 * 2.8346, 12 * 2.8346, f"第 {document.page} 页")
        canvas.restoreState()

    return draw


def _plain(value: str) -> str:
    value = re.sub(r"!\[([^\]]*)]\([^)]+\)", r"\1", value)
    value = re.sub(r"\[([^\]]+)]\([^)]+\)", r"\1", value)
    value = re.sub(r"[*_~`]+", "", value)
    return html.unescape(re.sub(r"<[^>]+>", "", value)).strip()


def _pdf_text(value: str) -> str:
    return html.escape(value, quote=False).replace("\n", "<br/>")


def _table_cells(line: str) -> list[str]:
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def _file_hash(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for block in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(block)
    return digest.hexdigest()


def _compatibility_report_path(report: Report) -> str:
    return f"reports/{report.id}/source.md"


def _json_dict(value: str | None) -> dict[str, Any] | None:
    if not value:
        return None
    try:
        parsed = json.loads(value)
    except json.JSONDecodeError:
        return None
    return parsed if isinstance(parsed, dict) else None


def _json_list(value: str | None) -> list[Any]:
    if not value:
        return []
    try:
        parsed = json.loads(value)
    except json.JSONDecodeError:
        return []
    return parsed if isinstance(parsed, list) else []
